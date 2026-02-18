# -*- coding: utf-8 -*-
"""
Génère le cahier des charges ASJ Espelette en .docx
Version 2.0 — Rédigé du point de vue du club (maître d'ouvrage)
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BLUE = RGBColor(0, 0x70, 0xA0)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF4, 0xF5, 0xF7)

def set_cell_bg(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shd)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '0070A0')
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_bg(cell, 'F4F5F7')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = BLUE
    hs.font.name = 'Calibri'
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)

# ===== PAGE DE TITRE =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CAHIER DES CHARGES')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Site Internet Officiel')
run.font.size = Pt(18)
run.font.color.rgb = DARK

club = doc.add_paragraph()
club.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = club.add_run('ASJ Espelette')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = BLUE

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    'Version 2.0 — Février 2026',
    'Émetteur : ASJ Espelette',
    'Rédigé avec l\'accompagnement de : SKenea',
    'Destinataire : Prestataire(s) consulté(s)',
    'Statut : Document de travail'
]:
    run = meta.add_run(line + '\n')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

doc.add_page_break()

# ===== NOTE INTRO =====
note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(12)
run = note.add_run('Ce document formalise les besoins de l\'ASJ Espelette en matière de site internet. '
    'Les points restant à arbitrer par la direction sont signalés par le picto [À PRÉCISER].')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = GRAY

# ===== 1. CONTEXTE ET OBJECTIFS =====
doc.add_heading('1. CONTEXTE ET OBJECTIFS', level=1)

doc.add_paragraph(
    'L\'ASJ Espelette est un club de football basé à Espelette (Pays Basque). '
    'À ce jour, le club ne dispose pas de site internet officiel. '
    'Un site Footeo existe (asjespelette.footeo.com), mais il concerne uniquement '
    'les U11 et n\'a pas été mis en place par la direction.')

doc.add_paragraph('Nous souhaitons nous doter d\'un site internet qui permette de :')
for item in [
    'Donner une vitrine officielle au club (pour les familles, les licenciés, les curieux)',
    'Centraliser les résultats et le calendrier des matchs',
    'Publier des actualités et des photos facilement',
    'Afficher nos partenaires et attirer de nouveaux sponsors',
    'Proposer le site en français et en euskara pour refléter l\'identité basque du club',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Ce que le site n\'a pas besoin de faire (sauf décision contraire de la direction) :', level=3)
for item in [
    'Gérer les inscriptions (c\'est le rôle de Sport Easy)',
    'Vendre des maillots ou de la billetterie',
    'Remplacer les réseaux sociaux (Instagram, Facebook)',
]:
    doc.add_paragraph(item, style='List Bullet')

# ===== 2. PÉRIMÈTRE FONCTIONNEL =====
doc.add_heading('2. PÉRIMÈTRE FONCTIONNEL', level=1)

doc.add_heading('Pages souhaitées', level=2)
add_styled_table(doc,
    ['Page', 'Contenu', 'Source des données'],
    [
        ['Accueil', 'Dernières actus, prochains matchs, liens réseaux sociaux', 'Bénévole / Automatique FFF'],
        ['Le Club', 'Histoire du club, bureau, photos des éducateurs', 'Fourni par le club'],
        ['Équipes', 'Seniors, U13 x3, U11, U9, U8, U7 + liens FFF et Sport Easy', 'Liens existants'],
        ['Calendrier & Résultats', 'Calendrier et scores de toutes les équipes', 'AUTOMATIQUE (données FFF)'],
        ['Galerie photos', 'Albums par événement (matchs, tournois, fêtes)', 'Bénévole (upload simple)'],
        ['Partenaires', 'Logos sponsors (Or/Argent/Bronze) + offre PDF', 'Fourni par le club'],
        ['Contact', 'Formulaire, adresse du stade, horaires, plan', '[À PRÉCISER]'],
        ['Mentions légales', 'Obligations RGPD, cookies, éditeur', 'À rédiger par le prestataire'],
    ],
    col_widths=[3.5, 7, 5]
)

doc.add_paragraph()
doc.add_heading('Fonctionnalités attendues', level=2)
for item in [
    'Interface d\'administration : nous souhaitons qu\'un bénévole non-technique puisse publier des actualités, '
    'ajouter des photos et gérer les partenaires sans toucher au code',
    'Intégration FFF : le calendrier des matchs et les résultats de toutes les équipes du club doivent '
    's\'afficher automatiquement, sans intervention humaine',
    'Bilinguisme français / euskara : tout le site doit être disponible dans les deux langues, '
    'avec un bouton de bascule visible en permanence',
    'Responsive : le site doit s\'afficher correctement sur mobile, tablette et ordinateur',
    'Formulaire de contact : les visiteurs doivent pouvoir envoyer un message au club',
    'Liens réseaux sociaux : les comptes Instagram et Facebook du club doivent être intégrés',
]:
    doc.add_paragraph(item, style='List Bullet')

# ===== 3. EXIGENCES ET CONTRAINTES =====
doc.add_heading('3. EXIGENCES ET CONTRAINTES', level=1)

doc.add_heading('Bilinguisme français / euskara', level=2)
doc.add_paragraph(
    'Tout le site devra être disponible en français et en euskara. '
    'Un bouton FR / EUS devra être visible en permanence en haut de chaque page. '
    'La langue par défaut sera le français.')
doc.add_paragraph(
    'Menus en basque : Hasiera, Kluba, Taldeak, Egutegia, Argazkiak, Bazkideak, Harremanetan jarri...')

p = doc.add_paragraph()
run = p.add_run('[À PRÉCISER] La traduction en euskara sera-t-elle assurée par un bénévole du club '
    'ou faut-il prévoir un traducteur externe ?')
run.bold = True
run.font.color.rgb = BLUE

doc.add_paragraph()
doc.add_heading('Données FFF (calendrier et résultats)', level=2)
doc.add_paragraph(
    'Nous souhaitons que les données suivantes soient récupérées automatiquement depuis le site de la '
    'Fédération Française de Football, sans aucune intervention de notre part :')
for item in [
    'Calendrier des matchs : dates, horaires, lieux, équipes adverses',
    'Résultats et scores : mis à jour après chaque journée',
    'Informations du club : catégories, terrains',
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('À noter : ')
run.bold = True
run.font.color.rgb = BLUE
p2 = doc.add_paragraph(
    'La FFF ne propose pas d\'API publique documentée pour les clubs amateurs. Il est néanmoins possible '
    'de récupérer les données via les interfaces internes du site fff.fr (mécanisme testé avec succès). '
    'Le prestataire devra mettre en place cette intégration et prévoir un suivi dans le cadre de la '
    'maintenance, car ces interfaces pourraient évoluer. Un mécanisme de secours et un lien direct vers '
    'la page du club sur fff.fr devront être prévus en dernier recours.')
for run2 in p2.runs:
    run2.italic = True
    run2.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_heading('Identité visuelle', level=2)
doc.add_paragraph(
    'Couleur officielle du club : bleu (référence FFF). '
    'Palette déduite du logo : bleu #0070A0. '
    'Nous fournirons notre logo en haute qualité (cf. section 4).')
p = doc.add_paragraph()
run = p.add_run('[À PRÉCISER] La palette de couleurs reste à valider par la direction.')
run.bold = True
run.font.color.rgb = BLUE

doc.add_paragraph()
doc.add_heading('Nom de domaine', level=2)
add_styled_table(doc,
    ['Élément', 'Choix envisagé'],
    [
        ['Nom de domaine', 'asjespelette.eus — l\'extension du Pays Basque'],
        ['Langue par défaut', 'Français, avec bascule vers euskara'],
    ],
    col_widths=[4, 11]
)

# ===== 4. CONTENUS FOURNIS PAR LE CLUB =====
doc.add_page_break()
doc.add_heading('4. CONTENUS FOURNIS PAR LE CLUB', level=1)
doc.add_paragraph(
    'Nous nous engageons à fournir les éléments suivants au prestataire retenu.')

add_styled_table(doc,
    ['#', 'Élément', 'Précision', 'Priorité'],
    [
        ['1', 'Logo du club en haute qualité', 'PNG fond transparent ou SVG. La version FFF est insuffisante.', 'Haute'],
        ['2', 'Texte de présentation / histoire', 'Quelques paragraphes. Le prestataire pourra nous accompagner.', 'Moyenne'],
        ['3', 'Photos des éducateurs', 'Une par coach, avec nom et catégorie. Format portrait.', 'Moyenne'],
        ['4', 'Président et composition du bureau', 'Pour la page Club et les mentions légales.', 'Haute'],
        ['5', 'Horaires d\'entraînement', 'Par catégorie, pour la page Contact.', 'Moyenne'],
        ['6', 'Numéro de téléphone du club', 'Pour la page Contact et les mentions légales.', 'Moyenne'],
        ['7', 'Liste des partenaires actuels', 'Nom, logo HD, site web, niveau (Or/Argent/Bronze).', 'Avant mise en ligne'],
        ['8', 'Contenu offre de partenariat', 'Avantages par niveau. Mis en forme en PDF par le prestataire.', 'Avant mise en ligne'],
        ['9', 'Comptes Instagram et Facebook', 'Les URLs exactes des comptes officiels du club.', 'Haute'],
    ],
    col_widths=[0.8, 4, 7, 2.5]
)

# ===== 5. DÉCISIONS À PRENDRE =====
doc.add_paragraph()
doc.add_heading('5. DÉCISIONS À PRENDRE PAR LA DIRECTION', level=1)
doc.add_paragraph(
    'Certains points restent à arbitrer en interne avant de lancer la réalisation.')
add_styled_table(doc,
    ['#', 'Question', 'Options envisagées'],
    [
        ['A', 'Sport Easy : quel niveau d\'intégration ?', 'Option 1 : Simple lien (suffisant pour commencer) / Option 2 : Afficher les infos directement (plus complexe)'],
        ['B', 'Boutique / billetterie : à prévoir ?', 'Option 1 : Pas pour le moment / Option 2 : Oui, dès le lancement'],
        ['C', 'Qui traduira en euskara ?', 'Un bénévole saisit les traductions / Ou prévoir un traducteur externe'],
        ['D', 'Qui administrera le site ?', 'Bénévole référent à désigner — il sera formé par le prestataire'],
    ],
    col_widths=[0.8, 5, 9]
)

# ===== 6. GESTION QUOTIDIENNE =====
doc.add_heading('6. GESTION QUOTIDIENNE DU SITE', level=1)
doc.add_paragraph(
    'Nous souhaitons que le site soit simple à administrer au quotidien. '
    'Il ne doit pas devenir une charge pour nos bénévoles.')

doc.add_heading('Ce qui doit être automatique (aucune intervention de notre part)', level=2)
for item in ['Calendrier et résultats : récupérés automatiquement depuis le site de la FFF',
             'Plan d\'accès : carte interactive générée automatiquement']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Ce que notre bénévole fera (via une interface simple)', level=2)
for item in [
    'Publier une actualité : remplir un formulaire (titre, texte, photo) — l\'article apparaît sur le site',
    'Ajouter des photos : glisser-déposer des images dans la galerie',
    'Mettre à jour les partenaires : opération ponctuelle (1 à 2 fois par an)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Ce que notre bénévole ne doit PAS avoir à faire', level=2)
for item in [
    'Toucher au code du site',
    'Gérer l\'hébergement ou la technique',
    'Mettre à jour des logiciels',
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Une formation est attendue pour le bénévole administrateur, '
    'ainsi qu\'un guide d\'utilisation simple.')
run.italic = True

# ===== 7. PRESTATIONS ATTENDUES =====
doc.add_heading('7. PRESTATIONS ATTENDUES DU PRESTATAIRE', level=1)
doc.add_paragraph('Le prestataire retenu devra fournir :')
for item in [
    'Conception et maquettes graphiques aux couleurs du club, validées par la direction avant développement',
    'Développement et intégration du site complet (toutes les pages listées en section 2)',
    'Configuration du bilinguisme français / euskara',
    'Intégration automatique du calendrier et des résultats depuis le site de la FFF',
    'Interface d\'administration permettant à un bénévole de gérer les actualités, les photos et les partenaires',
    'Formation du bénévole administrateur',
    'Guide d\'utilisation simple et illustré',
    'Mise en ligne du site sur le nom de domaine retenu',
    'Maintenance : une offre de maintenance annuelle devra être proposée, couvrant au minimum '
    'la surveillance du site, les sauvegardes, le suivi de l\'intégration FFF, et le support au bénévole',
]:
    doc.add_paragraph(item, style='List Bullet')

# ===== 8. SOLUTIONS TECHNIQUES =====
doc.add_page_break()
doc.add_heading('8. SOLUTIONS TECHNIQUES ENVISAGÉES', level=1)
doc.add_paragraph(
    'Trois grandes familles de solutions nous ont été présentées pour réaliser un site de ce type. '
    'Le choix définitif sera fait par la direction en fonction des priorités du club.')

# --- Option A ---
doc.add_heading('Option A — Site codé sur-mesure (HTML/CSS/JS + PHP)', level=2)
p = doc.add_paragraph()
run = p.add_run('Un site créé de zéro, taillé pour le club, avec une petite interface d\'administration.')
run.italic = True
run.font.color.rgb = GRAY

doc.add_heading('Avantages :', level=3)
for item in [
    'Site très rapide (pas de logiciel lourd)',
    'Aucune mise à jour de sécurité à gérer',
    'Coûts de fonctionnement les plus bas',
    'Design 100 % aux couleurs du club',
    'Le site appartient totalement au club',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Inconvénients :', level=3)
for item in [
    'Interface d\'administration plus basique',
    'Toute nouvelle fonctionnalité nécessite le prestataire',
    'Moins de flexibilité pour le bénévole sur la mise en page',
]:
    doc.add_paragraph(item, style='List Bullet')

# --- Option B ---
doc.add_paragraph()
doc.add_heading('Option B — WordPress', level=2)
p = doc.add_paragraph()
run = p.add_run('Le système de gestion de contenu le plus utilisé au monde (~43 % des sites web).')
run.italic = True
run.font.color.rgb = GRAY

doc.add_heading('Avantages :', level=3)
for item in [
    'Interface d\'administration riche et intuitive (éditeur visuel)',
    'Le bénévole peut modifier la mise en page, ajouter des pages',
    'Bilinguisme géré par un plugin (Polylang)',
    'Très grande communauté : facile de trouver de l\'aide',
    'Évolutif : boutique, newsletter, formulaires... via des plugins',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Inconvénients :', level=3)
for item in [
    'Mises à jour régulières nécessaires (sécurité)',
    'Sans maintenance, le site peut devenir vulnérable',
    'Un peu plus lent qu\'un site sur-mesure',
    'Coûts de fonctionnement un peu plus élevés (plugins, hébergement)',
]:
    doc.add_paragraph(item, style='List Bullet')

# --- Option C ---
doc.add_paragraph()
doc.add_heading('Option C — Wix', level=2)
p = doc.add_paragraph()
run = p.add_run('Un constructeur de sites en ligne, tout compris (hébergement inclus).')
run.italic = True
run.font.color.rgb = GRAY

doc.add_heading('Avantages :', level=3)
for item in [
    'Très facile à prendre en main (glisser-déposer)',
    'Hébergement et sécurité gérés par Wix',
    'Pas besoin de prestataire pour les modifications simples',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Inconvénients :', level=3)
for item in [
    'Abonnement mensuel obligatoire (~15-30 euros/mois)',
    'Bilinguisme limité',
    'Moins de contrôle sur le design et le code',
    'Dépendance totale à la plateforme Wix (pas de portabilité)',
    'Intégration FFF plus contrainte techniquement',
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('[À PRÉCISER] Le choix de la solution technique sera validé par la direction '
    'lors de la réunion de cadrage.')
run.bold = True
run.font.color.rgb = BLUE

# ===== 9. ESTIMATION BUDGÉTAIRE =====
doc.add_page_break()
doc.add_heading('9. ESTIMATION BUDGÉTAIRE', level=1)

doc.add_heading('Enveloppe de création : 3 000 euros HT', level=2)
doc.add_paragraph(
    'L\'enveloppe budgétaire envisagée pour la création du site est de 3 000 euros HT, couvrant :')
for item in [
    'Conception et maquettes graphiques',
    'Développement et intégration du site complet',
    'Configuration du bilinguisme français / euskara',
    'Intégration automatique des données FFF (calendrier et résultats)',
    'Interface d\'administration',
    'Formation du bénévole + guide d\'utilisation',
    'Mise en ligne',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Hébergement et nom de domaine (coûts récurrents à notre charge)', level=2)
add_styled_table(doc,
    ['Poste', 'Estimation annuelle'],
    [
        ['Nom de domaine (asjespelette.eus)', '~10-15 euros/an'],
        ['Hébergement', '~40-150 euros/an selon la solution retenue'],
        ['Total hébergement + domaine', '~50-165 euros/an'],
    ],
    col_widths=[8, 6]
)

doc.add_paragraph()
doc.add_heading('Maintenance annuelle', level=2)
doc.add_paragraph(
    'Nous souhaitons qu\'un contrat de maintenance annuel soit proposé pour assurer la pérennité '
    'du site. Il devra couvrir au minimum :')
for item in [
    'Surveillance du bon fonctionnement du site',
    'Sauvegardes régulières',
    'Suivi de l\'intégration FFF (ajustements si la Fédération modifie son site)',
    'Support au bénévole administrateur (questions, accompagnement)',
    'Corrections de bugs éventuels',
    'Mises à jour de sécurité si nécessaire',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
add_styled_table(doc,
    ['Poste', 'Estimation annuelle'],
    [
        ['Maintenance', '~250 euros HT/an'],
    ],
    col_widths=[8, 6]
)

doc.add_paragraph()
doc.add_heading('Synthèse sur 5 ans', level=2)
add_styled_table(doc,
    ['Poste', 'Estimation'],
    [
        ['Création du site', '~3 000 euros HT'],
        ['Hébergement + domaine (5 ans)', '~250-825 euros'],
        ['Maintenance (5 ans)', '~1 250 euros'],
        ['Total sur 5 ans', '~4 500-5 075 euros HT'],
    ],
    col_widths=[8, 6]
)

p = doc.add_paragraph()
run = p.add_run('Ces montants sont des estimations indicatives. '
    'Le détail sera précisé dans la proposition du prestataire retenu.')
run.italic = True
run.font.color.rgb = GRAY

# ===== 10. CALENDRIER =====
doc.add_heading('10. CALENDRIER PRÉVISIONNEL', level=1)
p = doc.add_paragraph()
run = p.add_run('Les délais ci-dessous s\'entendent après réception des contenus que nous fournirons '
    '(logo, textes, photos). La phase de recueil pourrait prendre 2 à 4 semaines.')
run.italic = True
run.font.color.rgb = GRAY

doc.add_paragraph()
add_styled_table(doc,
    ['Étape', 'Durée estimée', 'Qui fait quoi'],
    [
        ['Recueil des contenus', '2-4 sem.', 'Le club fournit les éléments (cf. section 4)'],
        ['Maquettes graphiques', '1 sem.', 'Prestataire propose, le club valide'],
        ['Développement et intégration', '2-4 sem.', 'Prestataire'],
        ['Configuration bilinguisme', '1 sem.', 'Prestataire + bénévole traducteur'],
        ['Tests et corrections', '1 sem.', 'Le club teste, prestataire corrige'],
        ['Formation du bénévole', '1h', 'Prestataire forme'],
        ['TOTAL après contenus', '5-8 semaines', ''],
    ],
    col_widths=[4.5, 2.5, 8]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Objectif de mise en ligne : septembre 2026 (rentrée sportive)')
run.bold = True
run.font.color.rgb = BLUE

# ===== 11. PROCHAINES ÉTAPES =====
doc.add_heading('11. PROCHAINES ÉTAPES', level=1)
steps = [
    'La direction relit ce document et arbitre les points [À PRÉCISER]',
    'Réunion de cadrage (30 min) pour répondre aux questions et valider les besoins',
    'Validation du budget et de l\'approche technique',
    'Le club rassemble les contenus (logo, textes, photos — cf. section 4)',
    'Démarrage de la réalisation dès que les contenus principaux sont transmis au prestataire',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.color.rgb = BLUE
    p.add_run(step)

# --- Footer ---
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document émis par l\'ASJ Espelette — rédigé avec l\'accompagnement de SKenea.')
run.font.size = Pt(9)
run.font.color.rgb = GRAY

# --- Save ---
output = r'c:\Users\Stephane (NoAdmin)\Documents\Stephane\Personnel\SKenea\SiteWebASJE\docs\CDC-ASJ-Espelette.docx'
doc.save(output)
print(f'Document généré : {output}')
